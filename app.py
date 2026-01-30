import os
import faiss
import numpy as np
import streamlit as st

from docx import Document
from PyPDF2 import PdfReader
from langchain_community.document_loaders import WebBaseLoader
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.docstore.in_memory import InMemoryDocstore
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_classic.text_splitter import RecursiveCharacterTextSplitter
from langchain_classic.chains import create_history_aware_retriever, create_retrieval_chain
from langchain_classic.chains.combine_documents import create_stuff_documents_chain
from langchain_core.documents import Document as LCDocument
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_core.messages import HumanMessage, AIMessage
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

# ============================================================================
# CONFIGURATION
# ============================================================================

# Get API keys from environment variables
huggingface_api_key = os.getenv('HUGGINGFACEHUB_API_TOKEN')
google_api_key = os.getenv('GOOGLE_API_KEY')

if not huggingface_api_key or not google_api_key:
    raise ValueError(
        "Missing API keys. Please set the following environment variables:\n"
        "  - HUGGINGFACEHUB_API_TOKEN\n"
        "  - GOOGLE_API_KEY\n"
        "\n"
        "You can set them by:\n"
        "  1. Creating a .env file in the project root (see .env.example)\n"
        "  2. Or setting them in your system environment variables"
    )

os.environ['HUGGINGFACEHUB_API_TOKEN'] = huggingface_api_key
os.environ['GOOGLE_API_KEY'] = google_api_key

# Text Splitting Configuration
CHUNK_SIZE = 500
CHUNK_OVERLAP = 200
CHUNK_SEPARATORS = ["\n\n", "\n", ".", " ", ""]

# Embedding Configuration
EMBEDDING_MODEL = "sentence-transformers/all-mpnet-base-v2"
EMBEDDING_DEVICE = 'cpu'

# LLM Configuration
LLM_MODEL = "gemini-3-pro-preview"
LLM_TEMPERATURE = 0.3
RETRIEVER_K = 5

# ============================================================================
# DOCUMENT PROCESSING
# ============================================================================

def create_text_splitter():
    """Create and return a configured text splitter."""
    return RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=CHUNK_SEPARATORS
    )


def process_web_links(urls, text_splitter):
    """Process web links and return documents."""
    documents = []
    for url in urls:
        loader = WebBaseLoader(url)
        raw_docs = loader.load()
        documents.extend(text_splitter.split_documents(raw_docs))
    return documents


def process_text_input(text, text_splitter):
    """Process plain text input and return documents."""
    texts = text_splitter.split_text(text)
    return [
        LCDocument(
            page_content=chunk,
            metadata={"source": "User Input Text"}
        )
        for chunk in texts
    ]


def extract_pdf_text(file):
    """Extract text from PDF file."""
    pdf_reader = PdfReader(file)
    return "\n".join(
        page.extract_text() 
        for page in pdf_reader.pages 
        if page.extract_text()
    )


def extract_docx_text(file):
    """Extract text from DOCX file."""
    doc = Document(file)
    return "\n".join(para.text for para in doc.paragraphs)


def extract_txt_text(file):
    """Extract text from TXT file."""
    return file.read().decode('utf-8')


def process_file(file, input_type, text_splitter):
    """Process a single file and return document chunks."""
    extractors = {
        "PDF": extract_pdf_text,
        "DOCX": extract_docx_text,
        "TXT": extract_txt_text
    }
    
    try:
        file_text = extractors[input_type](file)
        texts = text_splitter.split_text(file_text)
        
        return [
            LCDocument(
                page_content=chunk,
                metadata={"source": f"{file.name} (Chunk {i+1})"}
            )
            for i, chunk in enumerate(texts)
        ]
    except Exception as e:
        st.error(f"Error processing file {file.name}: {e}")
        return []


def process_files(files, input_type, text_splitter):
    """Process multiple files and return all documents."""
    documents = []
    for file in files:
        documents.extend(process_file(file, input_type, text_splitter))
    return documents


def load_documents(input_type, input_data):
    """Load and process documents based on input type."""
    text_splitter = create_text_splitter()
    
    if input_type == "Link":
        return process_web_links(input_data, text_splitter)
    elif input_type == "Text":
        return process_text_input(input_data, text_splitter)
    elif input_type in ["PDF", "DOCX", "TXT"]:
        return process_files(input_data, input_type, text_splitter)
    else:
        raise ValueError(f"Unsupported input type: {input_type}")


# ============================================================================
# VECTOR STORE CREATION
# ============================================================================

def create_embeddings():
    """Create and return HuggingFace embeddings model."""
    return HuggingFaceEmbeddings(
        model_name=EMBEDDING_MODEL,
        model_kwargs={'device': EMBEDDING_DEVICE},
        encode_kwargs={'normalize_embeddings': False}
    )


def create_vector_store(documents):
    """Create and populate a FAISS vector store from documents."""
    if not documents:
        raise ValueError("No valid text could be extracted from the inputs.")
    
    embeddings = create_embeddings()
    
    # Get embedding dimension
    sample_embedding = np.array(embeddings.embed_query("sample text"))
    dimension = sample_embedding.shape[0]
    
    # Create FAISS index
    index = faiss.IndexFlatL2(dimension)
    
    # Create vector store
    vector_store = FAISS(
        embedding_function=embeddings.embed_query,
        index=index,
        docstore=InMemoryDocstore(),
        index_to_docstore_id={},
    )
    
    vector_store.add_documents(documents)
    return vector_store


def process_input(input_type, input_data):
    """Main entry point for processing input and creating vector store."""
    documents = load_documents(input_type, input_data)
    return create_vector_store(documents)

# ============================================================================
# RAG CHAIN SETUP
# ============================================================================

def create_llm():
    """Create and return the LLM instance."""
    return ChatGoogleGenerativeAI(
        model=LLM_MODEL,
        temperature=LLM_TEMPERATURE,
        convert_system_message_to_human=True
    )


def create_history_aware_retriever_chain(llm, retriever):
    """Create history-aware retriever for contextual question reformulation."""
    system_prompt = (
        "Given a chat history and the latest user question "
        "which might reference context in the chat history, "
        "formulate a standalone question which can be understood "
        "without the chat history. Do NOT answer the question, "
        "just reformulate it if needed and otherwise return it as is."
    )
    
    prompt = ChatPromptTemplate.from_messages([
        ("system", system_prompt),
        MessagesPlaceholder("chat_history"),
        ("human", "{input}"),
    ])
    
    return create_history_aware_retriever(llm, retriever, prompt)


def create_qa_chain(llm):
    """Create question-answering chain with prompt."""
    system_prompt = (
        "You are an assistant for question-answering tasks. "
        "Use the following pieces of retrieved context to answer "
        "the question. If you don't know the answer, say that you "
        "don't know. "
        "Keep the answer concise but ensure it fully addresses the user's question."
        "\n\n"
        "{context}"
    )
    
    prompt = ChatPromptTemplate.from_messages([
        ("system", system_prompt),
        MessagesPlaceholder("chat_history"),
        ("human", "{input}"),
    ])
    
    return create_stuff_documents_chain(llm, prompt)


def get_rag_chain(vectorstore):
    """Create and return the complete RAG chain."""
    llm = create_llm()
    retriever = vectorstore.as_retriever(search_kwargs={"k": RETRIEVER_K})
    
    history_aware_retriever = create_history_aware_retriever_chain(llm, retriever)
    qa_chain = create_qa_chain(llm)
    
    return create_retrieval_chain(history_aware_retriever, qa_chain)


# ============================================================================
# CHAT HISTORY MANAGEMENT
# ============================================================================

def get_chat_history(messages):
    """Convert Streamlit messages to LangChain message format."""
    chat_history = []
    for msg in messages[:-1]:
        if msg["role"] == "user":
            chat_history.append(HumanMessage(content=msg["content"]))
        elif msg["role"] == "assistant":
            chat_history.append(AIMessage(content=msg["content"]))
    return chat_history

# ============================================================================
# STREAMLIT UI COMPONENTS
# ============================================================================

def initialize_session_state():
    """Initialize Streamlit session state variables."""
    if "vectorstore" not in st.session_state:
        st.session_state.vectorstore = None
    if "messages" not in st.session_state:
        st.session_state.messages = []


def render_sidebar_input():
    """Render sidebar input form and return input type and data."""
    st.header("Add Source File")
    input_type = st.selectbox("Input Type", ["Link", "PDF", "Text", "DOCX", "TXT"])
    input_data = None
    
    if input_type == "Link":
        number_input = st.number_input(
            min_value=1, max_value=20, step=1, 
            label="Number of Links"
        )
        input_data = [
            url for i in range(number_input)
            if (url := st.text_input(f"URL {i+1}"))
        ]
    elif input_type == "Text":
        input_data = st.text_area("Enter text content", height=200)
    elif input_type == "PDF":
        input_data = st.file_uploader(
            "Upload PDF(s)", 
            type=["pdf"], 
            accept_multiple_files=True
        )
    elif input_type == "TXT":
        input_data = st.file_uploader(
            "Upload Text file(s)", 
            type=['txt'], 
            accept_multiple_files=True
        )
    elif input_type == "DOCX":
        input_data = st.file_uploader(
            "Upload DOCX file(s)", 
            type=['docx', 'doc'], 
            accept_multiple_files=True
        )
    
    return input_type, input_data


def handle_data_processing(input_type, input_data):
    """Handle the data processing button click."""
    if not input_data:
        st.warning("Please provide input data.")
        return
    
    with st.spinner("Processing Knowledge Base..."):
        try:
            vectorstore = process_input(input_type, input_data)
            st.session_state.vectorstore = vectorstore
            st.session_state.messages = []
            
            count = len(input_data) if isinstance(input_data, list) else 1
            st.success(f"Knowledge Base Created from {count} file(s)!")
        except Exception as e:
            st.error(f"Error: {e}")


def render_chat_messages():
    """Render all chat messages in the conversation."""
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.write(message["content"])


def render_sources(source_documents):
    """Render source documents in an expandable section."""
    with st.expander("📚 View Sources"):
        for i, doc in enumerate(source_documents, 1):
            source_name = doc.metadata.get("source", "Unknown")
            st.markdown(f"**Source {i} ({source_name}):**")
            st.caption(doc.page_content[:300] + "...")


def handle_user_query(query):
    """Process user query and generate response."""
    st.session_state.messages.append({"role": "user", "content": query})
    
    with st.chat_message("assistant"):
        response_placeholder = st.empty()
        full_response = ""
        source_documents = []
        
        rag_chain = get_rag_chain(st.session_state.vectorstore)
        history = get_chat_history(st.session_state.messages)
        
        try:
            for chunk in rag_chain.stream({"input": query, "chat_history": history}):
                if "answer" in chunk:
                    full_response += chunk["answer"]
                    response_placeholder.markdown(full_response + "▌")
                if "context" in chunk:
                    source_documents = chunk["context"]
            
            response_placeholder.markdown(full_response)
            render_sources(source_documents)
            st.session_state.messages.append({"role": "assistant", "content": full_response})
        except Exception as e:
            st.error(f"Error during generation: {e}")


def render_sidebar():
    """Render the complete sidebar."""
    input_type, input_data = render_sidebar_input()
    
    if st.button("Process Data"):
        handle_data_processing(input_type, input_data)
    
    if st.session_state.vectorstore is not None:
        st.divider()
        if st.button("Clear Chat History"):
            st.session_state.messages = []
            st.rerun()


def render_main_area():
    """Render the main chat area."""
    if st.session_state.vectorstore is None:
        st.info("👈 Please upload a file or enter a link in the sidebar to start chatting.")
        return
    
    render_chat_messages()
    
    if query := st.chat_input("Submit your queries..."):
        with st.chat_message("user"):
            st.write(query)
        handle_user_query(query)


# ============================================================================
# MAIN APPLICATION
# ============================================================================

def main():
    """Main Streamlit application entry point."""
    st.set_page_config(page_title="RAG Q&A App", layout="wide")
    st.title("Insight Engine: Chat With Your Doc")
    
    initialize_session_state()
    
    with st.sidebar:
        render_sidebar()
    
    render_main_area()

if __name__ == "__main__":
    main()